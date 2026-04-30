"""Script to inspect the Devops_sample_Report.docx template structure."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document(r'C:\D_Drive\regime-platform\ppt\Devops_sample_Report.docx')

print("=== DOCUMENT SECTIONS ===")
for i, section in enumerate(doc.sections):
    print(f"Section {i}: page size {section.page_width.inches:.1f}x{section.page_height.inches:.1f} in, "
          f"margins: top={section.top_margin.inches:.2f}, bottom={section.bottom_margin.inches:.2f}, "
          f"left={section.left_margin.inches:.2f}, right={section.right_margin.inches:.2f}")

print("\n=== STYLES USED ===")
used_styles = set()
for para in doc.paragraphs:
    used_styles.add(para.style.name)
print(sorted(used_styles))

print("\n=== ALL PARAGRAPHS (first 120) ===")
for i, para in enumerate(doc.paragraphs[:120]):
    runs_info = []
    for r in para.runs:
        info = f"[bold={r.bold}, italic={r.italic}, size={r.font.size.pt if r.font.size else 'inherit'}, font={r.font.name}]"
        runs_info.append(f"'{r.text}'{info}")
    alignment = para.alignment
    align_str = {0: 'LEFT', 1: 'CENTER', 2: 'RIGHT', 3: 'JUSTIFY', None: 'None'}.get(alignment, str(alignment))
    print(f"[{i:03d}] Style='{para.style.name}' Align={align_str} | Text='{para.text[:100]}' | Runs={runs_info[:2]}")

print("\n=== TABLES ===")
for t_idx, table in enumerate(doc.tables):
    print(f"\nTable {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols")
    for r_idx, row in enumerate(table.rows[:5]):
        for c_idx, cell in enumerate(row.cells[:4]):
            print(f"  [{r_idx},{c_idx}]: '{cell.text[:60]}'")
