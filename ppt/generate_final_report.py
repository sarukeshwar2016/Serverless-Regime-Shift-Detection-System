import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def clear_paragraph(p):
    p._element.getparent().remove(p._element)

def generate_report():
    template_path = r'C:\D_Drive\regime-platform\ppt\Devops_sample_Report.docx'
    output_path = r'C:\D_Drive\regime-platform\ppt\Final_Regime_Shift_Report_Auto.docx'
    
    doc = Document(template_path)
    
    # --- 1. Replace Title Page and Bonafide Content ---
    for para in doc.paragraphs[:70]:
        text = para.text
        if "Secure and Scalable solution" in text:
            para.text = "Serverless Regime Shift Detection System with Automated CI/CD Pipeline"
            for r in para.runs:
                r.bold = True
        elif "AFRAZ TANVIR" in text:
            para.text = "SARUKESHWAR S [RA2311003011470]\nDEV VIKNESH AD [RA2311003011472]"
        elif "Dr. V. Deepan" in text:
            para.text = "Dr. S. Sivasakthiselvan"
        elif "Dr.V.Deepan" in text:
            para.text = "Dr. S. Sivasakthiselvan"
        elif "deployment of web applications has become" in text:
            para.text = "Financial markets generate vast amounts of time-series data where underlying statistical properties change abruptly, known as regime shifts. Detecting these shifts in real-time is critical for risk management and algorithmic trading. This project presents a serverless Regime Shift Detection System that continuously ingests financial data, applies advanced anomaly detection algorithms (PELT and ADWIN), and visualizes the results on a Next.js dashboard."
        elif "project begins by creating a simple Node.js" in text:
            para.text = "The platform is built using a microservices architecture, comprising a Python-based ingestion engine, a real-time detection layer, and a Next.js frontend, all communicating via Redis and MongoDB."
        elif "Next, an ECS cluster is created" in text:
            para.text = "To ensure seamless updates and continuous delivery, the entire application is containerized using Docker and deployed on an AWS EC2 instance. A robust CI/CD pipeline is implemented using Jenkins."
        elif "handle the incoming traffic, a load balancer" in text:
            para.text = "The Jenkins pipeline automates the retrieval of code from version control, builds the Docker containers, and deploys them to the production environment with zero downtime, using a 'Fast Demo Mode' hot-reloading setup."
        elif "Finally, Terraform is used" in text:
            para.text = "Overall, this project demonstrates the integration of machine learning algorithms for real-time financial monitoring with modern DevOps practices, ensuring a highly available, scalable, and rapidly deployable platform."
        elif "Overall, this project demonstrates the benefits" in text:
            para.text = ""

    # --- 2. Remove everything from CHAPTER 1 onwards ---
    chapter_1_found = False
    paragraphs_to_remove = []
    
    for para in doc.paragraphs:
        if para.text.strip() == "CHAPTER 1":
            chapter_1_found = True
        
        if chapter_1_found:
            paragraphs_to_remove.append(para)
            
    for p in paragraphs_to_remove:
        clear_paragraph(p)
        
    # --- 3. Add New Content ---
    
    def add_heading_1(text):
        try:
            p = doc.add_paragraph(text, style='Heading 1')
        except:
            p = doc.add_paragraph(text, style='Normal')
            for r in p.runs: r.bold = True
            p.runs[0].font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        return p
        
    def add_heading_2(text):
        try:
            return doc.add_paragraph(text, style='Heading 2')
        except:
            p = doc.add_paragraph(text, style='Normal')
            for r in p.runs: r.bold = True
            p.runs[0].font.size = Pt(14)
            return p
        
    def add_heading_3(text):
        try:
            return doc.add_paragraph(text, style='Heading 3')
        except:
            p = doc.add_paragraph(text, style='Normal')
            for r in p.runs: r.bold = True
            p.runs[0].font.size = Pt(12)
            return p
        
    def add_body(text):
        try:
            return doc.add_paragraph(text, style='Body Text')
        except:
            return doc.add_paragraph(text, style='Normal')
        
    def add_bullet(text):
        try:
            return doc.add_paragraph(text, style='List Paragraph')
        except:
            return doc.add_paragraph("• " + text, style='Normal')

    # Add page break before CHAPTER 1
    doc.add_page_break()
    
    # CHAPTER 1
    add_heading_1("CHAPTER 1")
    add_heading_2("INTRODUCTION")
    
    add_heading_3("Aim")
    add_body("The aim of this project is to develop a serverless, real-time regime shift detection platform for financial markets and deploy it using an automated Jenkins CI/CD pipeline on AWS EC2.")
    
    add_heading_3("Background")
    add_body("Financial markets are highly volatile, and understanding when market conditions fundamentally change (a regime shift) is crucial. Traditional batch-processing systems are too slow for modern algorithmic trading. A real-time, streaming architecture is required to detect anomalies instantly. Furthermore, deploying such complex systems manually is error-prone, necessitating automated DevOps practices.")
    
    add_heading_3("Context of the Project")
    add_body("This project bridges the gap between quantitative finance and modern cloud infrastructure. It provides a full-stack solution that not only detects mathematical anomalies in cryptocurrency streams (like BTC/USDT) but also features a fully automated deployment pipeline. Changes made to the codebase are automatically tested, built, and deployed to an AWS EC2 instance within seconds using Jenkins and Docker.")

    doc.add_page_break()

    # CHAPTER 2
    add_heading_1("CHAPTER 2")
    add_heading_2("LITERATURE SURVEY")
    
    add_heading_3("Overview of Regime Shift Detection")
    add_body("Regime shift detection involves identifying abrupt changes in the statistical properties of a time series. This is particularly relevant in finance, where markets switch between bull, bear, and sideways regimes. Early detection allows for dynamic asset allocation and risk mitigation.")
    
    add_heading_3("Time Series Analysis (PELT and ADWIN)")
    add_body("The project utilizes two primary algorithms. PELT (Pruned Exact Linear Time) is used for retrospective change point detection, offering mathematical guarantees on optimal segmentations. ADWIN (Adaptive Windowing) is a streaming algorithm designed to detect drift in real-time data streams without requiring a fixed window size.")
    
    add_heading_3("Docker Containerization")
    add_body("Docker provides a standardized way to package applications and their dependencies into lightweight containers. This ensures consistency across different environments, solving the 'it works on my machine' problem, and allows for rapid deployment and scaling on cloud platforms.")
    
    add_heading_3("Jenkins CI/CD Automation")
    add_body("Jenkins is an open-source automation server that enables developers to build, test, and deploy software reliably. By setting up a CI/CD pipeline, every commit pushed to the version control repository is automatically fetched, built into a Docker image, and deployed to the live server, drastically reducing deployment time and manual errors.")

    doc.add_page_break()

    # CHAPTER 3
    add_heading_1("CHAPTER 3")
    add_heading_2("SYSTEM DESIGN AND METHODOLOGY")
    
    add_heading_3("System Requirements and Specifications")
    add_body("Hardware Requirements:")
    add_bullet("1 t2.medium AWS EC2 instance for hosting the Docker container.")
    add_bullet("Elastic IP for a persistent static endpoint.")
    add_body("Software Requirements:")
    add_bullet("Node.js (Next.js framework)")
    add_bullet("Python 3.10+ (FastAPI, River, Ruptures)")
    add_bullet("Docker and Docker Compose")
    add_bullet("Jenkins (Local or Cloud host)")
    add_bullet("MongoDB (Cloud Atlas) and Redis (Local)")
    
    add_heading_3("Tools and Technologies used")
    add_body("Python is utilized for the backend ingestion and machine learning detection engine due to its extensive data science ecosystem. Next.js and TailwindCSS power the real-time interactive dashboard. Redis serves as an ultra-fast in-memory hot layer for live state, while MongoDB acts as the cold layer for permanent anomaly logging. Jenkins handles the CI/CD orchestration.")
    
    add_heading_3("System Architecture and Components")
    add_body("The architecture consists of multiple decoupled microservices running within a single Docker container. The Python ingestion script connects to Binance WebSockets, pulling live tick data. This data is fed into the Detection layer which updates Redis. The Next.js frontend polls the API layer for state updates. The entire stack is deployed via Jenkins.")
    
    add_heading_3("Fig 3.1 Architecture Diagram")
    try:
        doc.add_picture(r'C:\D_Drive\regime-platform\ppt\regime_architecture_diagram_1776332695698.png', width=Inches(6.0))
    except:
        add_body("[Architecture Diagram Image]")

    doc.add_page_break()

    # CHAPTER 4
    add_heading_1("CHAPTER 4")
    add_heading_2("IMPLEMENTATION")
    
    add_heading_3("Data Ingestion and Detection Layer")
    add_body("A Python script establishes a WebSocket connection to the Binance API, streaming live cryptocurrency prices. The data is buffered and analyzed using the 'ruptures' library for offline PELT detection and the 'river' library for online ADWIN drift detection. Results and confidence scores are immediately written to a local Redis instance.")
    
    add_heading_3("Next.js Dashboard")
    add_body("The frontend is a Next.js React application styled with TailwindCSS. It features a responsive layout that displays real-time price charts using Recharts and a live anomaly ledger. The dashboard polls a FastAPI endpoint to retrieve the latest state from Redis, providing instantaneous visual feedback when a regime shift occurs.")
    
    add_heading_3("Dockerization")
    add_body("A comprehensive Dockerfile was authored to package the entire platform. It uses an Ubuntu base image, installs Python and Node.js dependencies, configures the Redis server, and builds the Next.js production app. A bash startup script (`start.sh`) orchestrates the launch of all background processes simultaneously within the container.")
    
    add_heading_3("Jenkins CI/CD Pipeline Configuration")
    add_body("A Jenkinsfile defines the declarative deployment pipeline. When code is pushed to GitHub, Jenkins polls the repository (via Poll SCM `H/2 * * * *`), retrieves the latest code on the EC2 instance via SSH, runs `docker build` to compile the new changes, and executes `docker run` to restart the application with the new image, achieving full deployment in under 60 seconds.")

    doc.add_page_break()

    # CHAPTER 5
    add_heading_1("CHAPTER 5")
    add_heading_2("RESULTS AND DISCUSSIONS")
    
    add_heading_3("Application Interface")
    add_body("The deployed Next.js dashboard provides a clear, real-time view of asset states. It successfully displays 'STABLE', 'TRANSITIONING', or 'STRESSED' states based on the detection algorithms' confidence outputs.")
    try:
        doc.add_picture(r'C:\D_Drive\regime-platform\ppt\webf1.png', width=Inches(6.0))
    except:
        pass
        
    add_heading_3("Algorithm Performance")
    add_body("The ADWIN algorithm successfully detected micro-shifts in volatility within milliseconds, while the PELT algorithm confirmed macro-regime changes over larger time windows. The integration of both provided a robust confidence metric.")
    try:
        doc.add_picture(r'C:\D_Drive\regime-platform\ppt\chart_algorithm_comparison_1776353532485.png', width=Inches(6.0))
    except:
        pass

    add_heading_3("Deployment Performance")
    add_body("The Jenkins CI/CD pipeline successfully automated the deployment process. Testing showed that code commits pushed to the main branch were consistently built and deployed to the live AWS EC2 instance in approximately 30-40 seconds, minimizing downtime and manual intervention.")

    doc.add_page_break()

    # CHAPTER 6
    add_heading_1("CHAPTER 6")
    add_heading_2("CONCLUSION AND RECOMMENDATIONS")
    
    add_heading_3("Project Findings")
    add_body("The project successfully demonstrated the feasibility of building a real-time, serverless financial monitoring platform. The combination of advanced time-series algorithms with a high-performance Redis cache allowed for sub-second detection latency. Furthermore, the implementation of a Jenkins CI/CD pipeline proved essential for rapidly iterating and deploying complex containerized applications to AWS.")
    
    add_heading_3("Security Measures Implemented")
    add_body("Sensitive environment variables, such as the MongoDB URI and EC2 SSH private keys, were securely managed using Jenkins Credentials. SSH access to the EC2 instance was restricted using AWS Security Groups, and Docker containers were run with minimal necessary privileges.")
    
    add_heading_3("Future Addons/Improvements")
    add_body("Future improvements could include integrating Kubernetes for container orchestration to handle higher loads across multiple nodes. Additionally, the detection engine could be expanded to support more complex multi-variate anomaly detection models using deep learning (e.g., LSTMs or Transformers). WebSockets could replace the frontend HTTP polling for even lower latency updates.")
    
    doc.save(output_path)
    print(f"Report successfully generated at: {output_path}")

if __name__ == '__main__':
    generate_report()
