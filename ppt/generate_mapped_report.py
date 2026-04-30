import json
from docx import Document

def map_and_replace():
    template_path = r'C:\D_Drive\regime-platform\ppt\Devops_sample_Report.docx'
    output_path = r'C:\D_Drive\regime-platform\ppt\Regime_Shift_Final_Report.docx'
    
    doc = Document(template_path)
    
    # 1. Front Matter
    doc.paragraphs[0].text = "Serverless Regime Shift Detection System with Automated CI/CD Deployment"
    doc.paragraphs[6].text = "SARUKESHWAR S [RA2311003011470]\nDEV VIKNESH AD [RA2311003011472]"
    doc.paragraphs[9].text = "Dr. Sivasakthiselvan S"
    doc.paragraphs[50].text = "Dr. Sivasakthiselvan S"
    
    doc.paragraphs[37].text = "Certified that this project report \"Serverless Regime Shift Detection System with Automated CI/CD Deployment\" is the bonafide work of \"SARUKESHWAR S (RA2311003011470) and DEV VIKNESH AD (RA2311003011472)\" of III Year/VI Sem B.tech(CSE) who carried out the mini project work under my supervision."
    doc.paragraphs[38].text = ""
    
    # 2. Abstract
    doc.paragraphs[55].text = "Financial markets generate vast amounts of time-series data where underlying statistical properties change abruptly, known as regime shifts. Detecting these shifts in real-time is critical for algorithmic trading and risk management."
    doc.paragraphs[57].text = "The project begins by developing a sophisticated Python ingestion engine that connects to Binance WebSockets to stream cryptocurrency tick data. The data is processed using advanced algorithms like PELT and ADWIN to identify mathematical anomalies."
    doc.paragraphs[59].text = "Next, an ultra-fast Redis in-memory datastore is utilized as a hot layer, while MongoDB acts as a permanent cold layer for anomaly logging. A Next.js frontend polls the API to display the live regime states."
    doc.paragraphs[61].text = "To handle continuous delivery and seamless updates, the entire microservices architecture is containerized using Docker and deployed onto an AWS EC2 instance. This guarantees environment consistency and scalability."
    doc.paragraphs[63].text = "Finally, a Jenkins CI/CD pipeline is implemented to automate the deployment process. Any code pushed to the version control repository is automatically fetched, built, and deployed to the production server within seconds."
    doc.paragraphs[65].text = "Overall, this project demonstrates the integration of quantitative machine learning models with modern DevOps and CI/CD practices, resulting in a highly robust and scalable financial monitoring platform."

    # 3. Chapter 1: Introduction
    doc.paragraphs[113].text = "The aim of this project is to develop a serverless, real-time regime shift detection platform for financial markets and deploy it using an automated Jenkins CI/CD pipeline on AWS EC2."
    doc.paragraphs[117].text = "Financial markets are highly volatile, and understanding when market conditions fundamentally change is crucial. Deploying such complex streaming systems manually is error-prone, necessitating automated DevOps practices."
    doc.paragraphs[120].text = "This project bridges the gap between quantitative finance and modern cloud infrastructure. It provides a full-stack solution that detects mathematical anomalies and features a fully automated deployment pipeline."

    # 4. Chapter 2: Literature Survey
    doc.paragraphs[155].text = "Overview of Regime Shift Detection and Time Series Analysis:"
    doc.paragraphs[158].text = "Regime shift detection involves identifying abrupt changes in the statistical properties of a time series. This is particularly relevant in finance. Key algorithms include:"
    doc.paragraphs[159].text = "PELT (Pruned Exact Linear Time): Used for retrospective change point detection, offering mathematical guarantees."
    doc.paragraphs[160].text = "ADWIN (Adaptive Windowing): A streaming algorithm designed to detect drift in real-time data streams without requiring a fixed window size."
    doc.paragraphs[161].text = "Integration: Combining these provides both macro-regime and micro-volatility detection."
    doc.paragraphs[162].text = "Performance: Optimized to process high-frequency tick data with sub-second latency."
    
    doc.paragraphs[163].text = "Overview of Docker containerization and its advantages"
    doc.paragraphs[165].text = "Docker is an open-source platform that provides a standardized way to package applications and their dependencies."
    doc.paragraphs[166].text = "Portability: Docker containers can be run on any platform that supports Docker."
    doc.paragraphs[167].text = "Consistency: Docker containers provide a consistent runtime environment, eliminating 'it works on my machine' issues."
    doc.paragraphs[168].text = "Resource efficiency: Docker containers are lightweight and share the host OS kernel."
    
    doc.paragraphs[171].text = "Introduction to Jenkins CI/CD and Automation"
    doc.paragraphs[173].text = "Jenkins is an open-source automation server that enables developers to build, test, and deploy software reliably."
    doc.paragraphs[174].text = "Automation: Automatically triggers builds based on version control commits using Poll SCM."
    doc.paragraphs[175].text = "Speed: Drastically reduces deployment time from minutes to seconds."
    doc.paragraphs[176].text = "Reliability: Removes human error from the deployment process."
    
    doc.paragraphs[179].text = "Importance of AWS EC2 in Application Hosting"
    doc.paragraphs[182].text = "Amazon Elastic Compute Cloud (EC2) provides scalable computing capacity in the AWS cloud."
    doc.paragraphs[183].text = "Control: Provides complete control over the computing resources and network."
    doc.paragraphs[184].text = "Static Endpoints: Using an Elastic IP ensures a persistent address for web applications."

    # 5. Chapter 3: System Design
    doc.paragraphs[195].text = "To deploy the Regime Shift Detection Platform, the following system requirements are defined:"
    doc.paragraphs[198].text = "1 t2.medium AWS EC2 instance for hosting the Docker container."
    doc.paragraphs[199].text = "1 Elastic IP for persistent web access."
    doc.paragraphs[203].text = "Python 3.10+ and Node.js (Next.js framework)"
    doc.paragraphs[204].text = "Docker, Redis, and MongoDB"
    doc.paragraphs[205].text = "Jenkins CI/CD Automation Server"
    doc.paragraphs[206].text = "Git version control"
    
    doc.paragraphs[210].text = "The Next.js dashboard must be accessible via HTTP traffic over port 80."
    doc.paragraphs[211].text = "The EC2 instance must be configured with a security group that allows incoming traffic on port 80 and SSH on port 22."
    
    doc.paragraphs[214].text = "Tools and Technologies used: Python and Next.js"
    doc.paragraphs[215].text = "Python is utilized for the backend ingestion and machine learning detection engine. Next.js and TailwindCSS power the interactive dashboard."
    
    doc.paragraphs[217].text = "Redis and MongoDB"
    doc.paragraphs[219].text = "Redis serves as an ultra-fast in-memory hot layer for live state, while MongoDB acts as the cold layer for permanent anomaly logging."

    doc.paragraphs[220].text = "Jenkins"
    doc.paragraphs[222].text = "Jenkins handles the CI/CD orchestration, ensuring continuous deployment of code updates."

    doc.paragraphs[223].text = "Docker"
    doc.paragraphs[225].text = "Docker will be used to containerize the entire platform, combining the Python backend and Next.js frontend into a single deployable image."
    
    doc.paragraphs[245].text = "The system architecture consists of multiple decoupled microservices running within a single Docker container, deployed to an EC2 instance via Jenkins."
    doc.paragraphs[247].text = "Jenkins CI/CD Deployment Pipeline"
    doc.paragraphs[248].text = "A Jenkinsfile defines the declarative pipeline. When code is pushed to GitHub, Jenkins polls the repository, retrieves the latest code on the EC2 instance, runs docker build, and restarts the container."

    # 6. Chapter 4: Implementation
    doc.paragraphs[270].text = "Data Ingestion and Detection Layer"
    doc.paragraphs[271].text = "A Python script establishes a WebSocket connection to the Binance API, streaming live cryptocurrency prices. The data is buffered and analyzed using the 'ruptures' library for offline PELT detection and the 'river' library for online ADWIN drift detection."
    
    doc.paragraphs[336].text = "Next.js Dashboard Implementation"
    doc.paragraphs[337].text = "The frontend is a Next.js React application styled with TailwindCSS. It features a responsive layout that displays real-time price charts using Recharts and a live anomaly ledger. The dashboard polls a FastAPI endpoint to retrieve the latest state."
    
    doc.paragraphs[426].text = "Dockerization and Startup Scripts"
    doc.paragraphs[427].text = "A comprehensive Dockerfile was authored to package the entire platform. A bash startup script (start.sh) orchestrates the launch of all background processes simultaneously within the container."
    
    doc.paragraphs[448].text = "Jenkins Pipeline Configuration"
    doc.paragraphs[449].text = "The Jenkins pipeline automates the SSH connection to the EC2 server, pulling the latest git commits, and executing the Docker build commands to ensure the live application is always up to date."

    # 7. Chapter 5: Results
    doc.paragraphs[515].text = "The deployed Next.js dashboard provides a clear, real-time view of asset states. It successfully displays STABLE, TRANSITIONING, or STRESSED states."
    doc.paragraphs[518].text = "The ADWIN algorithm successfully detected micro-shifts in volatility within milliseconds, while the PELT algorithm confirmed macro-regime changes over larger time windows."
    
    doc.paragraphs[534].text = "The Jenkins CI/CD pipeline successfully automated the deployment process."
    doc.paragraphs[537].text = "Testing showed that code commits pushed to the main branch were consistently built and deployed to the live AWS EC2 instance in approximately 30-40 seconds."
    
    doc.paragraphs[565].text = "In summary, the combination of advanced time-series algorithms with a high-performance Redis cache allowed for sub-second detection latency."
    
    # 8. Chapter 6: Conclusion
    doc.paragraphs[579].text = "Sensitive environment variables, such as the MongoDB URI and EC2 SSH private keys, were securely managed using Jenkins Credentials."
    doc.paragraphs[581].text = "SSH access to the EC2 instance was restricted using AWS Security Groups, and Docker containers were run with minimal necessary privileges."
    
    doc.paragraphs[626].text = "The project successfully demonstrated the feasibility of building a real-time, serverless financial monitoring platform."
    doc.paragraphs[628].text = "Real-time accuracy: The dual-algorithm approach provided robust confidence metrics."
    doc.paragraphs[630].text = "Automated Deployment: Jenkins eliminated manual deployment errors."
    doc.paragraphs[632].text = "Scalable Architecture: Docker ensures the app can run consistently anywhere."
    
    doc.paragraphs[643].text = "There are several future addons and improvements that can be made to the system:"
    doc.paragraphs[646].text = "WebSockets: WebSockets could replace the frontend HTTP polling for even lower latency updates."
    doc.paragraphs[648].text = "Deep Learning: The detection engine could be expanded to support LSTMs or Transformers."
    doc.paragraphs[650].text = "Kubernetes: Container orchestration to handle higher loads across multiple nodes."

    # Replace specific abbreviations
    for para in doc.paragraphs[90:104]:
        if "ECS" in para.text:
            para.text = "EC2\tElastic Compute Cloud"
        if "ECR" in para.text:
            para.text = "CI/CD\tContinuous Integration / Continuous Deployment"
            
    # Update Tables if any refer to ECS/ECR
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if "ECS" in cell.text:
                    cell.text = cell.text.replace("ECS", "EC2")

    doc.save(output_path)
    print(f"Perfectly mapped report successfully generated at: {output_path}")

if __name__ == '__main__':
    map_and_replace()
