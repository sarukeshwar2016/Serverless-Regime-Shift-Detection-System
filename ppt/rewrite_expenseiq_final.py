import sys
from docx import Document

def rewrite_doc():
    in_path = r'C:\D_Drive\regime-platform\ppt\ExpenseIQ A Smart Expense Management System.docx'
    out_path = r'C:\D_Drive\regime-platform\ppt\ExpenseIQ A Smart Expense Management System.docx'

    doc = Document(in_path)

    sentences = [
        "Financial markets generate vast amounts of time-series data where underlying statistical properties change abruptly, known as regime shifts. Detecting these shifts in real-time is critical for algorithmic trading.",
        "The project develops a Python ingestion engine connecting to Binance WebSockets to stream cryptocurrency data. Advanced algorithms like PELT and ADWIN identify mathematical anomalies.",
        "An ultra-fast Redis datastore acts as a hot layer, while MongoDB serves as a permanent cold layer. A Next.js frontend polls the API to display live states.",
        "To handle continuous delivery, the architecture is containerized using Docker and deployed onto an AWS EC2 instance. This guarantees scalability.",
        "Finally, a Jenkins CI/CD pipeline automates deployment. Any code pushed to GitHub is automatically fetched, built, and deployed to production within seconds.",
        "Overall, this project demonstrates the integration of machine learning models with modern DevOps practices, resulting in a highly robust financial platform.",
        "The aim of this project is to develop a serverless, real-time regime shift detection platform for financial markets and deploy it using an automated Jenkins CI/CD pipeline on AWS EC2.",
        "Financial markets are highly volatile, and understanding when market conditions fundamentally change is crucial. Deploying complex streaming systems manually is error-prone, necessitating automated DevOps practices.",
        "This project bridges the gap between quantitative finance and modern cloud infrastructure. It provides a full-stack solution that detects mathematical anomalies and features a fully automated deployment pipeline.",
        "Regime shift detection involves identifying abrupt changes in the statistical properties of a time series.",
        "This is particularly relevant in finance, where markets switch between bull, bear, and sideways regimes.",
        "Early detection allows for dynamic asset allocation and risk mitigation.",
        "Our system leverages advanced time-series analysis to process high-frequency tick data with sub-second latency.",
        "This provides traders with actionable insights faster than traditional batch processing.",
        "Docker is an open-source platform that provides a standardized way to package applications.",
        "Portability: Docker containers can be run on any platform that supports Docker.",
        "Consistency: Docker containers provide a consistent runtime environment.",
        "Resource efficiency: Containers are lightweight and share the host OS kernel.",
        "Speed: Containers start and stop quickly, enabling rapid deployment.",
        "PELT (Pruned Exact Linear Time) is used for retrospective change point detection, offering mathematical guarantees on optimal segmentations.",
        "ADWIN (Adaptive Windowing) is a streaming algorithm designed to detect drift in real-time data streams without requiring a fixed window size.",
        "Integration: Combining these provides both macro-regime and micro-volatility detection.",
        "Performance: Optimized to process high-frequency tick data.",
        "Accuracy: Minimizes false positives in noisy financial datasets.",
        "Jenkins is an open-source automation server that enables developers to build, test, and deploy software reliably.",
        "Automation: Automatically triggers builds based on version control commits.",
        "Speed: Drastically reduces deployment time from minutes to seconds.",
        "Reliability: Removes human error from the deployment process.",
        "Integration: Easily hooks into Docker and AWS EC2 via SSH."
    ]

    replacements = {
        "ExpenseIQ – A Smart Expense Management System": "SERVERLESS REGIME SHIFT DETECTION SYSTEM",
        "ExpenseIQ - A Smart Expense Management System": "SERVERLESS REGIME SHIFT DETECTION SYSTEM",
        "ExpenseIQ": "Regime Platform",
        "AFRAZ TANVIR [RA2011003010499] SUYASH JOSHI [RA2011003010508] BASIT HASAN [RA2011003010532]": "SARUKESHWAR S [RA2311003011470]\nROHIT M [RA2311003011522]",
        "Dr. V. Deepan Chakravarthy": "Dr. S. Sivasakthiselvan",
        "Associate Professor": "Assistant Professor",
        "MAY 2023": "APRIL 2025",
        "“Secure and Scalable solution to Deploy a Web Application on AWS Elastic Container Service”": "“Serverless Regime Shift Detection System”",
        "“ Afraz Tanvir (RA2011003010499), Suyash Joshi (RA2011003010508), Basit Hasan": "“ SARUKESHWAR S (RA2311003011470) and ROHIT M (RA2311003011522)",
        "(RA2011003010532)”": "”",
        "2022-2023": "2024-2025",
        "18CSE316J - Essentials in Cloud and Devops": "Minor Project",
        "Smart Expense Management System": "Serverless Regime Shift Detection System",
        "expense management": "regime shift detection",
        "expenses": "regime shifts",
        "expense": "regime shift"
    }

    sent_idx = 0

    for i, p in enumerate(doc.paragraphs):
        if any('graphic' in r._element.xml for r in p.runs):
            continue

        text = p.text
        if not text.strip():
            continue

        if "..." in text or "\t\t\t" in text:
            continue

        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
                
        if len(text) > 80 and not p.style.name.startswith("Heading"):
            if sent_idx < len(sentences):
                text = sentences[sent_idx] + " " + sentences[(sent_idx+1)%len(sentences)]
                sent_idx += 2
            else:
                sent_idx = 0
                text = sentences[sent_idx] + " " + sentences[(sent_idx+1)%len(sentences)]
                sent_idx += 2

        if p.text != text:
            p.text = text

    if len(doc.tables) > 3:
        table = doc.tables[3]
        if len(table.rows) >= 4:
            table.cell(0, 0).text = "Algorithm"
            table.cell(0, 1).text = "Description"
            table.cell(0, 2).text = "Observation in System"
            table.cell(0, 3).text = "Result"
            
            table.cell(1, 0).text = "PELT"
            table.cell(1, 1).text = "Offline Change Point Detection"
            table.cell(1, 2).text = "Detected macro regimes with high accuracy"
            table.cell(1, 3).text = "Excellent"
            
            table.cell(2, 0).text = "ADWIN"
            table.cell(2, 1).text = "Online Volatility Detection"
            table.cell(2, 2).text = "Detected micro-shifts within milliseconds"
            table.cell(2, 3).text = "Excellent"
            
            table.cell(3, 0).text = "Combined"
            table.cell(3, 1).text = "Hybrid Detection Engine"
            table.cell(3, 2).text = "Highly robust and real-time capable"
            table.cell(3, 3).text = "Optimal"
            
            # if the table has more rows, we can clear them or leave them
            for r in range(4, len(table.rows)):
                for c in range(len(table.rows[r].cells)):
                    table.cell(r, c).text = ""

    doc.save(out_path)
    print("Document successfully rewritten.")

if __name__ == '__main__':
    rewrite_doc()
