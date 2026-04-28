
import sys
import io
from docx import Document
from docxcompose.composer import Composer

# Set encoding for output
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

def replace_text(doc, old_text, new_text):
    for para in doc.paragraphs:
        if old_text in para.text:
            # Simple replacement might lose formatting within runs
            # But let's try to be thorough
            para.text = para.text.replace(old_text, new_text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if old_text in para.text:
                        para.text = para.text.replace(old_text, new_text)

def process_report():
    template_path = r"C:\D_Drive\regime-platform\ppt\Project Report template.docx"
    source_path = r"C:\D_Drive\regime-platform\ppt\Regime_Shift_Detection_Report.docx"
    output_path = r"C:\D_Drive\regime-platform\ppt\Final_Regime_Shift_Detection_Report.docx"

    print("Loading documents...")
    template_doc = Document(template_path)
    source_doc = Document(source_path)

    # Define all possible placeholder variations
    replacements = {
        "PROJECT TITLE": "REGIME SHIFT DETECTION SYSTEM",
        "(GUIDE NAME )": "Dr. SIVASAKTHISELVAN S",
        "(GUIDE NAME)": "Dr. SIVASAKTHISELVAN S",
        "(Designation, Department)": "Assistant Professor, Computing Technologies",
        "DR. NAME OF THE GUIDE": "Dr. SIVASAKTHISELVAN S",
        "NAME OF THE GUIDE": "Dr. SIVASAKTHISELVAN S",
        "(DESIGNATION)": "Assistant Professor",
        "COMPUTATIONAL INTELLIGENCE": "COMPUTING TECHNOLOGIES",
        "with specialization in (SPECIALIZATION NAME)": "",
        "(SPECIALIZATION NAME)": "",
        "STUDENT1 NAME": "SARUKESHWAR S",
        "STUDENT2 NAME": "ROHIT M",
        "STUDENT 1": "SARUKESHWAR S",
        "STUDENT 2": "ROHIT M",
        "REG NO 1": "RA2311003011470",
        "REG NO 2": "RA2311003011522",
        "[REG NUM]": "RA2311003011470", # This will be tricky if there are two
    }

    print("Replacing placeholders...")
    # Order matters: replace specific long strings first
    
    # Custom replacement for Bonafide to handle multiple [REG NUM]
    for para in template_doc.paragraphs:
        if "STUDENT1 NAME [REG NUM], STUDENT2 NAME [REG NUM]" in para.text:
            para.text = para.text.replace("STUDENT1 NAME [REG NUM], STUDENT2 NAME [REG NUM]", "SARUKESHWAR S [RA2311003011470], ROHIT M [RA2311003011522]")
        
        # General replacements
        for old, new in replacements.items():
            if old in para.text:
                para.text = para.text.replace(old, new)

    for table in template_doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for old, new in replacements.items():
                        if old in para.text:
                            para.text = para.text.replace(old, new)

    # Extract Abstract
    source_abstract = ""
    found_abstract = False
    for para in source_doc.paragraphs:
        if "ABSTRACT" in para.text.upper():
            found_abstract = True
            continue
        if found_abstract:
            if para.text.strip().upper().startswith("CHAPTER") or para.text.strip().upper().startswith("TABLE OF CONTENTS"):
                break
            source_abstract += para.text + "\n"
    
    if source_abstract:
        for para in template_doc.paragraphs:
            if "Provide a summary of the project" in para.text:
                para.text = source_abstract.strip()
                break

    # Save temp front matter
    print("Preparing front matter...")
    # Delete chapters
    cp_idx = -1
    for i, p in enumerate(template_doc.paragraphs):
        if "CHAPTER 1" in p.text.upper():
            cp_idx = i
            break
    
    if cp_idx != -1:
        p_list = template_doc.paragraphs
        for i in range(len(p_list)-1, cp_idx-1, -1):
            p = p_list[i]._element
            p.getparent().remove(p)
    
    # Prepare source body
    print("Preparing body...")
    src_cp_idx = -1
    for i, p in enumerate(source_doc.paragraphs):
        if "CHAPTER 1" in p.text.upper():
            src_cp_idx = i
            break
    
    if src_cp_idx != -1:
        p_list = source_doc.paragraphs
        for i in range(src_cp_idx - 1, -1, -1):
            p = p_list[i]._element
            p.getparent().remove(p)
    
    temp_front = "temp_front.docx"
    temp_body = "temp_body.docx"
    template_doc.save(temp_front)
    source_doc.save(temp_body)

    # Merge
    print("Merging...")
    master = Document(temp_front)
    composer = Composer(master)
    composer.append(Document(temp_body))
    composer.save(output_path)
    print(f"Success! Final report saved to {output_path}")

if __name__ == "__main__":
    process_report()
